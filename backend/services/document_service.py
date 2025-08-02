"""
Doküman İşleme Servisi
=======================
"""

import os
import uuid
from typing import Optional, List
from fastapi import HTTPException, status, UploadFile
from sqlalchemy.orm import Session
from app.config import settings
from models.document import Document, DocumentCreate, DocumentUpdate
from models.user import User
from services.ai_service import AIService

class DocumentService:
    """Doküman işleme servisi"""
    
    @staticmethod
    def save_upload_file(upload_file: UploadFile, user_id: int) -> tuple[str, str]:
        """Dosyayı kaydet ve yol döndür"""
        # Dosya uzantısını kontrol et
        file_extension = os.path.splitext(upload_file.filename)[1].lower()
        if file_extension not in settings.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Desteklenmeyen dosya türü: {file_extension}"
            )
        
        # Benzersiz dosya adı oluştur
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
        
        # Dosyayı kaydet
        try:
            with open(file_path, "wb") as buffer:
                content = upload_file.file.read()
                if len(content) > settings.MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="Dosya boyutu çok büyük"
                    )
                buffer.write(content)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Dosya kaydedilemedi: {str(e)}"
            )
        
        return file_path, unique_filename
    
    @staticmethod
    def extract_text_from_file(file_path: str, file_type: str) -> str:
        """Dosyadan metin çıkar"""
        try:
            if file_type == "pdf":
                return DocumentService._extract_from_pdf(file_path)
            elif file_type == "docx":
                return DocumentService._extract_from_docx(file_path)
            elif file_type == "txt":
                return DocumentService._extract_from_txt(file_path)
            else:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Desteklenmeyen dosya türü: {file_type}"
                )
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Metin çıkarılamadı: {str(e)}"
            )
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """PDF'den metin çıkar"""
        try:
            import PyPDF2
            text = ""
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"PDF işleme hatası: {str(e)}")
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """DOCX'den metin çıkar"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"DOCX işleme hatası: {str(e)}")
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """TXT'den metin çıkar"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as file:
                    return file.read().strip()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise Exception(f"TXT işleme hatası: {str(e)}")
        
        # Son çare olarak binary olarak oku ve decode et
        try:
            with open(file_path, "rb") as file:
                content = file.read()
                return content.decode('utf-8', errors='ignore').strip()
        except Exception as e:
            raise Exception(f"TXT işleme hatası: {str(e)}")
    
    @staticmethod
    def create_document(db: Session, user: User, upload_file: UploadFile, title: str = None) -> Document:
        """Yeni doküman oluştur"""
        # Dosyayı kaydet
        file_path, filename = DocumentService.save_upload_file(upload_file, user.id)
        
        # Dosya türünü belirle
        file_extension = os.path.splitext(upload_file.filename)[1].lower()
        file_type = file_extension[1:]  # .pdf -> pdf
        
        # Metin çıkar
        content = DocumentService.extract_text_from_file(file_path, file_type)
        
        # Başlık belirle
        document_title = title if title else upload_file.filename
        
        # Doküman oluştur
        document_data = DocumentCreate(
            title=document_title,
            filename=filename,
            file_size=upload_file.size,
            file_type=file_type
        )
        
        db_document = Document(
            title=document_data.title,
            filename=document_data.filename,
            file_path=file_path,
            file_size=document_data.file_size,
            file_type=document_data.file_type,
            content=content,
            user_id=user.id
        )
        
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        
        return db_document
    
    @staticmethod
    def get_user_documents(db: Session, user_id: int, skip: int = 0, limit: int = 100) -> List[Document]:
        """Kullanıcının dokümanlarını getir"""
        return db.query(Document).filter(Document.user_id == user_id).offset(skip).limit(limit).all()
    
    @staticmethod
    def get_document(db: Session, document_id: int, user_id: int) -> Optional[Document]:
        """Belirli bir dokümanı getir"""
        return db.query(Document).filter(
            Document.id == document_id,
            Document.user_id == user_id
        ).first()
    
    @staticmethod
    def delete_document(db: Session, document_id: int, user_id: int) -> bool:
        """Dokümanı sil"""
        document = DocumentService.get_document(db, document_id, user_id)
        if not document:
            return False
        
        # Dosyayı fiziksel olarak sil
        try:
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
        except Exception:
            pass  # Dosya silinmezse devam et
        
        db.delete(document)
        db.commit()
        return True
    
    @staticmethod
    async def process_document_with_ai(db: Session, document_id: int, user_id: int) -> Document:
        """Dokümanı AI ile işle"""
        document = DocumentService.get_document(db, document_id, user_id)
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Doküman bulunamadı"
            )
        
        if not document.content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Doküman içeriği bulunamadı"
            )
        
        # AI ile özet ve anahtar kelimeler oluştur
        try:
            summary, keywords = await AIService.generate_summary_and_keywords(document.content)
            
            # Dokümanı güncelle
            document.summary = summary
            document.keywords = keywords
            db.commit()
            db.refresh(document)
            
            return document
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"AI işleme hatası: {str(e)}"
            ) 